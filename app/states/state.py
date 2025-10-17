import reflex as rx
import os
import google.generativeai as genai
import logging
from typing import Literal
from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import random
import PyPDF2
from PIL import Image

GenerationStatus = Literal["Ready", "Generating...", "Complete", "Error"]
logging.basicConfig(level=logging.INFO)


class State(rx.State):
    """The main state for the application."""

    old_resume: str = ""
    job_description: str = ""
    status: GenerationStatus = "Ready"
    generated_content: str | None = None
    generated_file_path: str | None = None
    error_message: str = ""
    uploaded_file_name: str | None = None
    file_size: int = 0
    upload_status: str = "No file selected."

    def _extract_text_from_file(self, file_path: str) -> str:
        """Extracts text from various file types."""
        file_extension = os.path.splitext(file_path)[1].lower()
        try:
            if file_extension == ".pdf":
                with open(file_path, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    return """
""".join((page.extract_text() for page in reader.pages))
            elif file_extension in [".docx"]:
                doc = DocxDocument(file_path)
                return """
""".join([para.text for para in doc.paragraphs])
            elif file_extension == ".txt":
                with open(file_path, "r", encoding="utf-8") as f:
                    return f.read()
            elif file_extension in [".jpg", ".jpeg", ".png"]:
                image = Image.open(file_path)
                return f"[Image detected: {os.path.basename(file_path)} ({image.width}x{image.height}) - OCR not implemented]"
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
        except Exception as e:
            logging.exception(f"Failed to extract text from {file_path}: {e}")
            return f"[Error extracting text from file: {os.path.basename(file_path)}]"

    @rx.event
    async def handle_upload(self, files: list[rx.UploadFile]):
        """Handles the file upload and text extraction."""
        if not files:
            return
        file = files[0]
        upload_data = await file.read()
        upload_dir = rx.get_upload_dir()
        upload_dir.mkdir(parents=True, exist_ok=True)
        file_path = upload_dir / file.name
        with file_path.open("wb") as f:
            f.write(upload_data)
        self.uploaded_file_name = file.name
        self.file_size = len(upload_data)
        self.upload_status = "Extracting text..."
        try:
            extracted_text = self._extract_text_from_file(file_path)
            self.old_resume = extracted_text
            self.upload_status = "File processed successfully."
            yield rx.toast.success("File uploaded and text extracted.")
        except Exception as e:
            logging.exception(f"Failed to process file: {e}")
            self.upload_status = f"Error: {e}"
            self.old_resume = ""
            yield rx.toast.error(f"Failed to process file: {e}")

    def _add_paragraph_border(
        self, paragraph, position="bottom", size="4", color="auto", space="1"
    ):
        """Adds a border to the bottom of a paragraph."""
        p_pr = paragraph._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
        border_el = OxmlElement(f"w:{position}")
        border_el.set(qn("w:val"), "single")
        border_el.set(qn("w:sz"), size)
        border_el.set(qn("w:space"), space)
        border_el.set(qn("w:color"), color)
        p_bdr.append(border_el)

    def _create_docx(self, content: str) -> str:
        """Creates a professionally formatted .docx file from the generated content."""
        document = DocxDocument()
        for section in document.sections:
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
        style = document.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)
        font.color.rgb = RGBColor(0, 0, 0)
        lines = content.strip().split("""
""")
        if not lines:
            document.add_paragraph("No content generated.")
        else:
            name_p = document.add_paragraph()
            name_run = name_p.add_run(lines[0].strip())
            name_run.font.name = "Calibri"
            name_run.font.size = Pt(18)
            name_run.bold = True
            name_run.font.color.rgb = RGBColor(31, 78, 121)
            name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            name_p.paragraph_format.space_after = Pt(2)
            if len(lines) > 1:
                contact_p = document.add_paragraph()
                contact_run = contact_p.add_run(lines[1].strip())
                contact_run.font.size = Pt(10)
                contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                contact_p.paragraph_format.space_after = Pt(8)
                self._add_paragraph_border(contact_p, color="B0B0B0")
            current_section = None
            for line in lines[2:]:
                line_stripped = line.strip()
                if not line_stripped:
                    continue
                if line_stripped.isupper() and len(line_stripped.split()) <= 4:
                    if current_section is not None:
                        document.add_paragraph().paragraph_format.space_before = Pt(4)
                    heading = document.add_paragraph()
                    run = heading.add_run(line_stripped)
                    run.font.name = "Calibri"
                    run.bold = True
                    run.font.size = Pt(13)
                    run.font.color.rgb = RGBColor(31, 78, 121)
                    heading.paragraph_format.space_after = Pt(6)
                    self._add_paragraph_border(heading, size="6", color="D9D9D9")
                    current_section = line_stripped
                    continue
                if (
                    line_stripped.startswith("*")
                    or line_stripped.startswith("-")
                    or line_stripped.startswith("•")
                ):
                    p = document.add_paragraph(
                        line_stripped[1:].strip(), style="List Bullet"
                    )
                    p.paragraph_format.left_indent = Inches(0.35)
                    p.paragraph_format.space_after = Pt(4)
                    p.paragraph_format.line_spacing = 1.15
                    continue
                date_pattern = "(\\b(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\\s+\\d{4}\\s*[-–—]\\s*(?:Present|Current|(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\\s+\\d{4}))"
                date_match = re.search(date_pattern, line_stripped, re.IGNORECASE)
                if date_match and current_section in [
                    "PROFESSIONAL EXPERIENCE",
                    "EXPERIENCE",
                ]:
                    date_str = date_match.group(1)
                    title_str = (
                        line_stripped[: date_match.start()].strip().rstrip(" |,-")
                    )
                    p = document.add_paragraph()
                    p.paragraph_format.space_before = Pt(8)
                    p.paragraph_format.space_after = Pt(2)
                    title_run = p.add_run(title_str)
                    title_run.bold = True
                    title_run.font.size = Pt(11.5)
                    p.add_run("\t")
                    date_run = p.add_run(date_str)
                    date_run.font.italic = True
                    date_run.font.size = Pt(10)
                    date_run.font.color.rgb = RGBColor(89, 89, 89)
                    tab_stops = p.paragraph_format.tab_stops
                    tab_stops.add_tab_stop(
                        Inches(6.75), alignment=WD_ALIGN_PARAGRAPH.RIGHT
                    )
                    continue
                p = document.add_paragraph(line_stripped)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.15
        upload_dir = rx.get_upload_dir()
        upload_dir.mkdir(parents=True, exist_ok=True)
        filename = f"ats_resume_{random.randint(1000, 9999)}.docx"
        file_path = upload_dir / filename
        document.save(file_path)
        return filename

    @rx.var
    def is_generating(self) -> bool:
        """Check if the resume is currently being generated."""
        return self.status == "Generating..."

    @rx.var
    def generation_complete(self) -> bool:
        """Check if the resume generation is complete and successful."""
        return self.status == "Complete" and self.generated_file_path is not None

    @rx.event(background=True)
    async def handle_generation(self):
        """
        Handles the resume generation logic by calling the Gemini API.
        This is a background task to avoid blocking the UI.
        """
        async with self:
            if not self.old_resume or not self.job_description:
                self.status = "Error"
                self.error_message = "Resume and Job Description cannot be empty."
                yield rx.toast.error(self.error_message)
                return
            self.error_message = ""
            self.status = "Generating..."
            self.generated_content = None
            self.generated_file_path = None
        try:
            api_key = os.getenv("GEMINI_API_KEY")
            if not api_key:
                raise ValueError("GEMINI_API_KEY environment variable not set.")
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel("gemini-2.5-flash")
            prompt = f"\n                You are an expert ATS (Applicant Tracking System) resume writer. Your task is to rewrite the provided **Old Resume Text** to perfectly match the keywords, skills, and requirements found in the **Job Description**. The regenerated resume must be highly optimized, clear, and professional. Output the final resume in a **structured, clean plain-text format** with clear section headings (e.g., CONTACT, SUMMARY, EXPERIENCE, EDUCATION, SKILLS) for easy post-processing.\n\n                **Old Resume Text:**\n                {self.old_resume}\n\n                **Job Description:**\n                {self.job_description}\n                "
            response = await model.generate_content_async(prompt)
            async with self:
                self.generated_content = response.text
                logging.info("Successfully generated resume content.")
                filename = self._create_docx(self.generated_content)
                self.generated_file_path = filename
                self.status = "Complete"
                yield rx.toast.success("Resume generated successfully!")
        except Exception as e:
            logging.exception(f"Error during resume generation: {e}")
            async with self:
                self.status = "Error"
                self.error_message = f"An error occurred: {str(e)}"
            yield rx.toast.error(
                f"Error: {self.error_message}", duration=8000, close_button=True
            )